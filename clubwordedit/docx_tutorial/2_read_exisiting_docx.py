# 读取docx内部文字信息

# from docx import Document
#
# doc = Document('model.docx')
#
# for p in doc.paragraphs:
#     print(p.text)


#读取docx内部表格信息
from docx import Document

# 获取文档
doc = Document("model.docx")
# 获取 Word 文档中的所有表格
tables = doc.tables
# 选择第一个表格
table = tables[0]
values = []

# 遍历表格的每一行
for row in table.rows:
    # 遍历每一行中的单元格
    for cell in row.cells:
        # 将单元格中的数据添加到list中
        values.append(cell.text)
    value = ' '.join(values)
    print(value)
    values = []
