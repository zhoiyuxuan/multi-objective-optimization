# #添加文字
# from docx import Document
# # 写入文档
# doc = Document()
# # 添加标题
# doc.add_heading("一级标题", level=1)
# # 添加段落
# p2 = doc.add_paragraph("第二个段落")
# # 将新段落添加到已经有的段落之前
# p1 = p2.insert_paragraph_before("第一个段落")
#
# p3 = doc.add_paragraph("新段落")
# # 追加内容
# p3.add_run("加粗").bold = True
# p3.add_run("以及")
# p3.add_run("斜体").italic = True
#
# doc.save("new_doc.docx")



# #添加图片
# from docx import Document
# from docx.shared import Inches
#
# doc = Document()
# # 添加图片
# doc.add_picture("1.jpg", width=Inches(1.25))
# doc.save("new_pic.docx")



# 添加表格
from docx import Document

doc = Document()
# 创建table
table = doc.add_table(rows=3, cols=4)
# 设置table样式
table.style = "Table Grid"

# 第一种方法 先获取行 再获取该行中对应的单元格
row = table.rows[0]
row.cells[0].text = "第一行第一列"

# 第二种方法 直接指行号和列号
cell = table.cell(0, 1)
cell.text = "第一行第一列"

doc.save('newtable.docx')
